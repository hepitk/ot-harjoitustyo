import unittest
from services.document_handler import document_handler
from docx import Document
from entities.replace_data import ReplaceData


class TestDocumentHandler(unittest.TestCase):
    def setUp(self):
        self.filename = "asiakirjapohjat/viranhaltijapäätös_määräalan_myynti_pohja"
        self.document = Document(self.filename + ".docx")
        self.document_filled = Document(
            self.document.save("valmiit asiakirjat/valmis.docx"))
        self.user_question = "Syötä kiinteistötunnus:"
        self.placeholder = "[kiinteistötunnus]"
        self.replace_word = "[kana]"
        self.replace_data = ReplaceData(
            self.filename, self.user_question, self.placeholder)
    
    def test_set_font_works(self):
        document_handler.set_font("Arial", 12)
        self.assertEqual(document_handler.font_name, "Arial")
        self.assertEqual(document_handler.font_size, 12)

    def test_replace_word_works(self):
        replace_amount = document_handler.replace_words(
            self.document, self.replace_word, self.placeholder)
        self.document_filled = Document("valmiit asiakirjat/valmis.docx")
        for par in self.document_filled.paragraphs:
            if self.replace_word in par.text and replace_amount == 2:
                self.assertEqual(par.text, par.text)
                return True
        self.assertEqual("False", "True")

    
