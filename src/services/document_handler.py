from docx.shared import Pt
from docx import Document


class DocumentHandler:
    """Luokka, joka vastaa asiakirjapohjien muokkaamisesta ja valmiiden asiakirjojen luomisesta."""
    def __init__(self):
        """Luokan konstruktori.

        Args:
            font_name: Merkkijono, joka kuvaa fontin nimeä.
            font_size: Merkkijono, joka kuvaa fontin kokoa.
        """
        self.font_name = "Calibri"
        self.font_size = 11

    def set_font(self, font_name, font_size):
        """Muuttaa sanojen korvauksessa käytettävän fontin nimen ja koon.

        Args:
            font_name: Merkkijono, joka kuvaa fontin nimeä.
            font_size: Kokonaisluku, joka kuvaa fontin kokoa.
        Returns:
            Palauttaa True, kun metodi on suoritettu.
        """
        self.font_name = font_name
        self.font_size = font_size
        return True

    def get_font_name (self):
        """Palauttaa tämänhetkisen fontin nimen.

        Returns:
            Palauttaa fontin nimen merkkijonona.
        """
        return self.font_name
    
    def get_font_size (self):
        """Palauttaa tämänhetkisen fontin koon

        Returns:
            Palauttaa fontin koon kokonaislukuna.
        """
        return self.font_size

    def replace_words(self, document, user_input, placeholder):
        """Korvaa paikkatiedon halutulla käyttäjän syötteellä, ja luo valmiin asiakirjan.

        Args:
            document: Muokattava asiakirjapohja docx-kirjaston Document-muodossa
            user_input: Merkkijono, joka kuvaa käyttäjän syötettä ja jolla halutaan korvata paikkatietomerkintä asiakirjapohjassa.
            placeholder: Merkkijono, joka kuvaa paikkatietomerkintää asiakirjapohjassa.
        Returns:
            Palauttaa korvattujen sanojen lukumäärän.
        """
        print (type(self.font_size), flush=True)
        replace_amount = 0

        style = document.styles["Normal"]
        font = style.font
        font.name = self.font_name
        font.size = Pt(self.font_size)
        section = document.sections[0]
        header = section.header
        footer = section.footer

        for par in document.paragraphs:
            if placeholder in par.text:
                replace_amount += 1
                par.text = par.text.replace(placeholder, user_input)
        for par in header.paragraphs:
            if placeholder in par.text:
                replace_amount += 1
                par.text = par.text.replace(placeholder, user_input)
        for par in footer.paragraphs:
            if placeholder in par.text:
                replace_amount += 1
                par.text = par.text.replace(placeholder, user_input)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        if placeholder in par.text:
                            replace_amount += 1
                            par.text = par.text.replace(placeholder, user_input)        

        document.save("valmiit asiakirjat/valmis.docx")
        return replace_amount


document_handler = DocumentHandler()
